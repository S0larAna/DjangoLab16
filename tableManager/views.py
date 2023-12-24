from django.shortcuts import render, redirect
from django.http import HttpResponse
from .forms import *
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from .decorators import *
from django.contrib.auth.models import Group
from django.views.decorators.csrf import csrf_exempt

from django import template
from django.contrib.auth.models import Group

from openpyxl import Workbook
from docx import Document
import aspose.pdf as ap

def export_to_pdf_jewelry(request):
    # Get data from your model (adjust the queryset as needed)
    queryset = Jewelry.objects.all()

    # Create a PDF document
    pdf_document = ap.Document()
    page = pdf_document.pages.add()

    # Set up table layout

    # Add headers
    headers = ['author', 'type', 'material', 'defects', 'date', 'price', 'store', 'client']
    for col_num, header in enumerate(headers):
        text_fragment = ap.text.TextFragment(header)
        page.paragraphs.add(text_fragment)

    # Add data rows
    for obj in queryset:
        for col_num, field_name in enumerate(headers):
            value = str(getattr(obj, field_name, ''))
            text_fragment = ap.text.TextFragment("Hello,world!")
            page.paragraphs.add(text_fragment)

    # Create response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=exported_data.pdf'
    pdf_document.save(response)

    return response

def export_to_docx_jewelry(request):
    # Get data from your model (adjust the queryset as needed)
    queryset = Jewelry.objects.all()

    # Create a new Word document
    document = Document()

    # Write headers
    headers = ['author', 'type', 'material', 'defects', 'date', 'price', 'store', 'client']  # Replace with actual field names
    table = document.add_table(rows=1, cols=len(headers))
    for col_num, header in enumerate(headers):
        table.cell(0, col_num).text = header

    # Write data rows
    for obj in queryset:
        row_cells = table.add_row().cells
        for col_num, field_name in enumerate(headers):
            value = str(getattr(obj, field_name, ''))
            row_cells[col_num].text = value

    # Create response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=exported_data.docx'
    document.save(response)

    return response

def export_to_xls_jewelry(request):
    # Get data from your model (adjust the queryset as needed)
    queryset = Jewelry.objects.all()

    # Create a workbook and add a worksheet
    workbook = Workbook()
    sheet = workbook.active

    headers = ['author', 'type', 'material', 'defects', 'date', 'price', 'store', 'client']  # Replace with actual field names
    for col_num, header in enumerate(headers, 1):
        sheet.cell(row=1, column=col_num, value=header)

    # Write data rows
    for row_num, obj in enumerate(queryset, 2):
        for col_num, field_name in enumerate(headers, 1):
            value = getattr(obj, field_name)
            if isinstance(value, Store) or isinstance(value, Client) or isinstance(value, User):
                value = str(value)
            sheet.cell(row=row_num, column=col_num, value=value)

    # Create response
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename=exported_data.xls'
    workbook.save(response)

    return response

@login_required(login_url='login')
@admin_only
def home(request):
    stores = Store.objects.all()
    clients = Client.objects.all()
    return render(request, 'dashboard.html', {'stores': stores, 'clients': clients})

@login_required(login_url='login')
@allowed_users(allowed_roles='admin')
def jewelry(request):
    items = Jewelry.objects.all()
    return render(request, 'jewelry.html', {'items':items})

@login_required(login_url='login')
@allowed_users(allowed_roles='customer')
def userJewelry(request):
    items = Jewelry.objects.filter(author=request.user)
    print(Jewelry.objects.all()[0].author)
    return render(request, 'jewelry.html', {'items':items})

@login_required(login_url='login')
@allowed_users(allowed_roles='customer')
@csrf_exempt
def createUserJewelry(request):
    form = JewelryForm()
    if request.method=='POST':
        form = JewelryForm(request.POST)
        if form.is_valid():
            form.instance.author = request.user
            form.save()
            return redirect('user_jewelry')
        else:
            print('Form errors:', form.errors)
    context = {'form': form}
    return render(request, 'jewelry_form.html', context)

@login_required(login_url='login')
@allowed_users(allowed_roles='customer')
def deleteUserJewelry(request, pk):
    jewelry = Jewelry.objects.get(id=pk)
    if request.method == 'POST':
        jewelry.delete()
        return redirect('user_jewelry')
    context = {'item': jewelry}
    return render(request, 'delete_jewelry.html', context)

@login_required(login_url='login')
@allowed_users(allowed_roles='customer')
@csrf_exempt
def updateUserJewelry(request, pk):
    jewelry = Jewelry.objects.get(id=pk)
    form = JewelryForm(instance=jewelry)
    if request.method=='POST':
        form = JewelryForm(request.POST, instance=jewelry)
        if form.is_valid():
            form.instance.author = request.user
            form.save()
            return redirect('user_jewelry')
    context ={'form': form}
    return render(request, 'jewelry_form.html', context)


@login_required(login_url='login')
@allowed_users(allowed_roles='admin')
@csrf_exempt
def createJewelry(request):
    form = JewelryForm()
    if request.method=='POST':
        form = JewelryForm(request.POST)
        if form.is_valid():
            form.instance.author = request.user
            form.save()
            return redirect('/jewelry')
    context = {'form': form}
    return render(request, 'jewelry_form.html', context)


@login_required(login_url='login')
@allowed_users(allowed_roles='admin')
@csrf_exempt
def updateJewelry(request, pk):
    jewelry = Jewelry.objects.get(id=pk)
    form = JewelryForm(instance=jewelry)
    if request.method=='POST':
        form = JewelryForm(request.POST, instance=jewelry)
        if form.is_valid():
            form.instance.author = request.user
            form.save()
            return redirect('/jewelry')
    context ={'form': form}
    return render(request, 'jewelry_form.html', context)


@login_required(login_url='login')
@allowed_users(allowed_roles='admin')
def deleteJewelry(request, pk):
    jewelry = Jewelry.objects.get(id=pk)
    if request.method == 'POST':
        jewelry.delete()
        return redirect('/jewelry')
    context = {'item': jewelry}
    return render(request, 'delete_jewelry.html', context)


@login_required(login_url='login')
@allowed_users(allowed_roles='admin')
def clients(request):
    clients = Client.objects.all()
    return render(request, 'clients.html', {'clients': clients})


@login_required(login_url='login')
@allowed_users(allowed_roles='admin')
@csrf_exempt
def createClient(request):
    form = ClientForm()
    if request.method=='POST':
        form = ClientForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('/clients')

    context = {'form': form}
    return render(request, 'clients_form.html', context)


@login_required(login_url='login')
@allowed_users(allowed_roles='admin')
@csrf_exempt
def updateClient(request, pk):
    client = Client.objects.get(id=pk)
    form = ClientForm(instance=client)
    if request.method=='POST':
        form = ClientForm(request.POST, instance=client)
        if form.is_valid():
            form.save()
            return redirect('/clients')
    context ={'form': form}
    return render(request, 'clients_form.html', context)


@login_required(login_url='login')
@allowed_users(allowed_roles='admin')
def deleteClient(request, pk):
    client = Client.objects.get(id=pk)
    if request.method == 'POST':
        client.delete()
        return redirect('/clients')
    context = {'item': client}
    return render(request, 'delete_client.html', context)


@login_required(login_url='login')
@allowed_users(allowed_roles='admin')
@csrf_exempt
def stores(request):
    stores = Store.objects.all()
    return render(request, 'stores.html', {'stores': stores})


@login_required(login_url='login')
@allowed_users(allowed_roles='admin')
@csrf_exempt
def createStore(request):
    form = StoreForm()
    if request.method=='POST':
        form = StoreForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('/stores')

    context = {'form': form}
    return render(request, 'store_form.html', context)


@login_required(login_url='login')
@allowed_users(allowed_roles='admin')
@csrf_exempt
def updateStore(request, pk):
    store = Store.objects.get(id=pk)
    form = StoreForm(instance=store)
    if request.method=='POST':
        form = StoreForm(request.POST, instance=store)
        if form.is_valid():
            form.save()
            return redirect('/stores')
    context ={'form': form}
    return render(request, 'store_form.html', context)


@login_required(login_url='login')
@allowed_users(allowed_roles='admin')
def deleteStore(request, pk):
    store = Store.objects.get(id=pk)
    if request.method == 'POST':
        store.delete()
        return redirect('/stores')
    context = {'item': store}
    return render(request, 'delete_store.html', context)


@login_required(login_url='login')
def profile(request):
    return HttpResponse('Profile')

@unauthenticated_user
@csrf_exempt
def register(request):
    form = CreateUserForm()
    if request.method=='POST':
        form = CreateUserForm(request.POST)
        if form.is_valid():
            user = form.save()
            group = Group.objects.get(name="customer")
            user.groups.add(group)
            messages.success(request, 'Account successfully created')
            return redirect('login')
    context = {'form': form}
    return render(request, 'register.html', context)


@unauthenticated_user
@csrf_exempt
def loginUser(request):
    if request.method=='POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('dashboard')
        else:
            messages.info(request, "Username or password is incorrect")
    context={}
    return render(request, 'login.html', context)


@login_required(login_url='login')
def logoutUser(request):
    logout(request)
    return redirect('login')

def user_page(request):
    return render(request, 'user.html')

